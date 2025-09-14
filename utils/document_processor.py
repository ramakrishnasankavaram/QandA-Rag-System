import os
import PyPDF2
import docx
from typing import List, Dict
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f"Error reading PDF {file_path}: {str(e)}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {str(e)}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            print(f"Error reading TXT {file_path}: {str(e)}")
        return text
    
    def process_document(self, file_path: str) -> List[Document]:
        """Process a document and return chunked documents."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        if not text.strip():
            raise ValueError(f"No text extracted from {file_path}")
        
        # Create document with metadata
        document = Document(
            page_content=text,
            metadata={
                "source": file_path,
                "filename": os.path.basename(file_path)
            }
        )
        
        # Split document into chunks
        chunks = self.text_splitter.split_documents([document])
        
        return chunks
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[Document]:
        """Process multiple documents and return all chunks."""
        all_chunks = []
        for file_path in file_paths:
            try:
                chunks = self.process_document(file_path)
                all_chunks.extend(chunks)
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                continue
        return all_chunks